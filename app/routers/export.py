import io
import re

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

router = APIRouter(prefix="/export", tags=["export"])


class ExportDocxRequest(BaseModel):
    html: str
    nombre: str = "documento"


def _set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)


def _apply_font(run, bold=False, italic=False, underline=False, highlight=False):
    run.font.name      = "Times New Roman"
    run.font.size      = Pt(12)
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.underline = underline
    if highlight:
        run.font.highlight_color = 4  # WD_COLOR_INDEX.YELLOW


def _add_inline(para, node, bold=False, italic=False, underline=False):
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                run = para.add_run(text)
                _apply_font(run, bold=bold, italic=italic, underline=underline)
        elif isinstance(child, Tag):
            b = bold or child.name in ("strong", "b")
            i = italic or child.name in ("em", "i")
            u = underline or child.name == "u"
            if child.name == "mark":
                run = para.add_run(child.get_text())
                _apply_font(run, bold=b, italic=i, underline=u, highlight=True)
            else:
                _add_inline(para, child, bold=b, italic=i, underline=u)


def _process(doc: Document, el) -> None:
    if isinstance(el, NavigableString):
        return

    tag = el.name
    if tag is None:
        return

    if tag in ("h1", "h2", "h3"):
        level = int(tag[1])
        para  = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(el.get_text())
        run.font.name = "Times New Roman"
        run.font.size = Pt(16 - level * 2)
        run.font.bold = True

    elif tag == "p":
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_inline(para, el)

    elif tag == "ul":
        for li in el.find_all("li", recursive=False):
            para = doc.add_paragraph(style="List Bullet")
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_inline(para, li)

    elif tag == "ol":
        for li in el.find_all("li", recursive=False):
            para = doc.add_paragraph(style="List Number")
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_inline(para, li)

    elif tag in ("div", "body", "section", "article"):
        for child in el.children:
            _process(doc, child)

    else:
        text = el.get_text().strip()
        if text:
            para = doc.add_paragraph()
            run  = para.add_run(text)
            _apply_font(run)


def _html_to_docx(html: str) -> Document:
    doc = Document()
    _set_margins(doc)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    soup = BeautifulSoup(html, "html.parser")
    for child in soup.children:
        _process(doc, child)

    return doc


@router.post("/docx")
async def export_docx(body: ExportDocxRequest):
    doc    = _html_to_docx(body.html)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    nombre = re.sub(r"[^\w\-]", "_", body.nombre)[:60]
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{nombre}.docx"'},
    )
