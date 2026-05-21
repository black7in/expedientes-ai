import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def exportar_a_docx(contenido: str, nombre_archivo: str = "documento") -> bytes:
    """
    Convierte el texto generado a un .docx con formato legal boliviano.
    Detecta títulos (líneas en MAYÚSCULAS terminadas en .-) y negritas (**texto**).
    """
    doc = Document()

    # Márgenes legales: 2.5cm superior/inferior, 3cm izquierdo, 2.5cm derecho
    for section in doc.sections:
        section.top_margin    = Inches(0.98)   # ~2.5cm
        section.bottom_margin = Inches(0.98)
        section.left_margin   = Inches(1.18)   # ~3cm
        section.right_margin  = Inches(0.98)

    # Fuente base
    estilo = doc.styles["Normal"]
    fuente = estilo.font
    fuente.name = "Times New Roman"
    fuente.size = Pt(12)

    parrafos = contenido.split("\n\n")

    for bloque in parrafos:
        bloque = bloque.strip()
        if not bloque:
            continue

        # Título de sección: línea corta toda en mayúsculas terminada en .-
        if _es_titulo(bloque):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(bloque)
            run.bold = True
        else:
            # Párrafo normal con posible negritas inline
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _agregar_con_negritas(p, bloque)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Helpers ───────────────────────────────────────────────────────────────────

def _es_titulo(texto: str) -> bool:
    """Detecta líneas que son títulos de sección legal."""
    lineas = texto.split("\n")
    if len(lineas) > 2:
        return False
    primera = lineas[0].strip()
    return (
        primera == primera.upper()
        and len(primera) < 80
        and (primera.endswith(".-") or primera.endswith("-"))
    )


def _agregar_con_negritas(parrafo, texto: str) -> None:
    """Parsea **negrita** inline y agrega runs al párrafo."""
    partes = re.split(r"(\*\*[^*]+\*\*)", texto)
    for parte in partes:
        if parte.startswith("**") and parte.endswith("**"):
            run = parrafo.add_run(parte[2:-2])
            run.bold = True
        else:
            # Saltos de línea simples dentro del bloque → espacio
            parte_limpia = parte.replace("\n", " ")
            parrafo.add_run(parte_limpia)
