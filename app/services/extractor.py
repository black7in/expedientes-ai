import os
import tempfile
import fitz  # PyMuPDF
import mammoth
from docx import Document


def extraer_pdf(ruta: str) -> dict:
    """Extrae texto de un archivo PDF. Operación CPU-bound — invocar via asyncio.to_thread."""
    doc = fitz.open(ruta)
    try:
        partes: list[str] = []
        for pagina in doc:
            texto = pagina.get_text("text").strip()
            if texto:
                partes.append(texto)

        metadata = {k: v for k, v in (doc.metadata or {}).items() if v}
        num_paginas = doc.page_count
    finally:
        doc.close()

    return {
        "texto_completo": "\n\n".join(partes),
        "paginas":        num_paginas,
        "metadata":       metadata,
    }


def extraer_docx(ruta: str) -> dict:
    """Extrae texto de un archivo DOCX, incluyendo texto dentro de tablas."""
    doc = Document(ruta)
    partes: list[str] = []

    # Párrafos a nivel de cuerpo (fuera de tablas)
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            partes.append(t)

    # Celdas de tablas — evita duplicados por celdas fusionadas
    seen: set[int] = set()
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                cell_id = id(celda._tc)
                if cell_id in seen:
                    continue
                seen.add(cell_id)
                t = celda.text.strip()
                if t:
                    partes.append(t)

    texto_completo = "\n\n".join(partes)

    return {
        "texto_completo": texto_completo,
        "paginas": 1,
        "parrafos": len(partes),
        "metadata": {},
    }


def extraer_txt(ruta: str) -> dict:
    """Extrae texto de un archivo .txt."""
    with open(ruta, encoding="utf-8", errors="replace") as f:
        texto = f.read()
    return {"texto_completo": texto, "paginas": 1, "metadata": {}}


def extraer_html_docx(ruta: str) -> dict:
    """Extrae HTML semántico de un DOCX preservando negritas, cursivas y encabezados."""
    with open(ruta, "rb") as f:
        result = mammoth.convert_to_html(f)
    return {"html": result.value, "paginas": 1, "metadata": {}}


def extraer_html_pdf(ruta: str) -> dict:
    """Extrae HTML semántico de un PDF usando PyMuPDF."""
    doc = fitz.open(ruta)
    try:
        partes: list[str] = []
        for pagina in doc:
            html_pagina = pagina.get_text("html")
            if html_pagina.strip():
                partes.append(html_pagina)
        num_paginas = doc.page_count
        metadata = {k: v for k, v in (doc.metadata or {}).items() if v}
    finally:
        doc.close()

    return {
        "html":     "\n".join(partes),
        "paginas":  num_paginas,
        "metadata": metadata,
    }


def extraer_html_txt(ruta: str) -> dict:
    """Envuelve texto plano en párrafos HTML básicos."""
    with open(ruta, encoding="utf-8", errors="replace") as f:
        texto = f.read()
    parrafos = [f"<p>{p.strip()}</p>" for p in texto.split("\n\n") if p.strip()]
    return {"html": "\n".join(parrafos), "paginas": 1, "metadata": {}}


def extraer_html(ruta: str, formato: str) -> dict:
    """Extrae HTML semántico según el formato del archivo."""
    if not os.path.exists(ruta):
        raise FileNotFoundError(f"Archivo no encontrado: {ruta}")

    if formato == "pdf":
        return extraer_html_pdf(ruta)
    elif formato == "docx":
        return extraer_html_docx(ruta)
    elif formato == "txt":
        return extraer_html_txt(ruta)
    else:
        raise ValueError(f"Formato no soportado para extracción HTML: {formato}")


def extraer_texto(ruta: str, formato: str) -> dict:
    """Dispatcher: extrae texto según el formato del archivo."""
    if not os.path.exists(ruta):
        raise FileNotFoundError(f"Archivo no encontrado: {ruta}")

    if formato == "pdf":
        return extraer_pdf(ruta)
    elif formato == "docx":
        return extraer_docx(ruta)
    elif formato == "txt":
        return extraer_txt(ruta)
    else:
        raise ValueError(f"Formato no soportado: {formato}")
